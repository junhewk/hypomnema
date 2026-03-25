"""File parsing: PDF, DOCX, Markdown text extraction and storage."""

from __future__ import annotations

import dataclasses
import logging
import re
import tempfile
from pathlib import Path
from typing import TYPE_CHECKING

if TYPE_CHECKING:
    import aiosqlite

from docx import Document as DocxDocument
from pypdf import PdfReader

from hypomnema.db.models import Document

logger = logging.getLogger(__name__)

_MIME_MAP: dict[str, str] = {
    ".pdf": "application/pdf",
    ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ".md": "text/markdown",
}


class UnsupportedFormatError(ValueError):
    """Raised when a file extension is not supported."""


@dataclasses.dataclass(frozen=True)
class ParsedFile:
    """Result of text extraction from a file."""

    text: str
    mime_type: str
    title: str | None = None


@dataclasses.dataclass(frozen=True)
class ParsedPdf:
    """Text and page metadata extracted from a PDF."""

    text: str
    page_count: int


_PAGE_NUMBER_RE = re.compile(r"^(?:page\s+)?\d+(?:\s*(?:/|of)\s*\d+)?$", re.IGNORECASE)
_LIST_OR_HEADING_RE = re.compile(
    r"^(?:"
    r"[-*+•]\s+"
    r"|\d+[.)]\s+"
    r"|\d+(?:\.\d+)+\s+"
    r"|(?:figure|table)\s+\d+[.:]?\s+"
    r"|(?:abstract|keywords?)\b"
    r"|(?:references|bibliography|appendix|acknowledg(?:e)?ments?)\b"
    r")",
    re.IGNORECASE,
)
_SENTENCE_END_RE = re.compile(r'[.!?]["”’)\]]?$')
_MARGIN_REPEAT_MIN_PAGES = 3
_MARGIN_LINE_MAX_LEN = 160
_BACKMATTER_HEADING_RE = re.compile(
    r"^(?:references|bibliography|appendix(?:\b|[\s:.-])|appendices(?:\b|[\s:.-])|acknowledg(?:e)?ments?)",
    re.IGNORECASE,
)
_BACKMATTER_MIN_BLOCKS = 24
_BACKMATTER_MIN_FRACTION = 0.35


def _trim_pdf_backmatter(blocks: list[str]) -> list[str]:
    if len(blocks) < _BACKMATTER_MIN_BLOCKS:
        return blocks

    min_index = max(8, int(len(blocks) * _BACKMATTER_MIN_FRACTION))
    for index, block in enumerate(blocks):
        if index < min_index:
            continue
        if _BACKMATTER_HEADING_RE.match(block.strip()):
            return blocks[:index]
    return blocks


def _normalize_margin_line(line: str) -> str:
    normalized = line.strip().casefold()
    normalized = re.sub(r"\bpage\s+\d+(?:\s+of\s+\d+)?\b", " ", normalized)
    normalized = re.sub(r"\b\d+\b", " ", normalized)
    normalized = re.sub(r"\s+", " ", normalized)
    return normalized.strip()


def _is_probable_margin_line(line: str) -> bool:
    stripped = line.strip()
    if not stripped:
        return False
    if len(stripped) > _MARGIN_LINE_MAX_LEN:
        return False
    if _PAGE_NUMBER_RE.fullmatch(stripped):
        return True
    normalized = _normalize_margin_line(stripped)
    return bool(normalized) and len(normalized) <= _MARGIN_LINE_MAX_LEN


def _detect_repeated_margin_lines(page_lines: list[list[str]]) -> set[str]:
    if len(page_lines) < _MARGIN_REPEAT_MIN_PAGES:
        return set()

    counts: dict[str, int] = {}
    for lines in page_lines:
        candidates = [line for line in lines if _is_probable_margin_line(line)]
        window = candidates[:2] + candidates[-2:]
        seen_for_page: set[str] = set()
        for line in window:
            normalized = _normalize_margin_line(line)
            if not normalized or normalized in seen_for_page:
                continue
            seen_for_page.add(normalized)
            counts[normalized] = counts.get(normalized, 0) + 1

    min_pages = max(_MARGIN_REPEAT_MIN_PAGES, len(page_lines) // 2)
    return {line for line, count in counts.items() if count >= min_pages}


def _is_structural_pdf_line(line: str) -> bool:
    stripped = line.strip()
    if not stripped:
        return False
    if _LIST_OR_HEADING_RE.match(stripped):
        return True
    if stripped.isupper() and len(stripped) <= 120:
        return True
    if re.match(r"^\[\d+\]", stripped):
        return True
    return bool(stripped.lower().startswith("doi:"))


def _should_start_new_pdf_paragraph(current: str, next_line: str) -> bool:
    if not current:
        return False
    if _is_structural_pdf_line(current) or _is_structural_pdf_line(next_line):
        return True
    return bool(_SENTENCE_END_RE.search(current.strip()))


def _coalesce_pdf_lines(lines: list[str]) -> list[str]:
    paragraphs: list[str] = []
    current = ""

    for raw_line in lines:
        line = raw_line.strip()
        if not line:
            if current:
                paragraphs.append(current.strip())
                current = ""
            continue
        if _PAGE_NUMBER_RE.fullmatch(line):
            continue
        if not current:
            current = line
            continue
        if current.endswith("-") and line[:1].islower():
            current = current[:-1] + line
            continue
        if _should_start_new_pdf_paragraph(current, line):
            paragraphs.append(current.strip())
            current = line
            continue
        current = f"{current} {line}"

    if current:
        paragraphs.append(current.strip())
    return paragraphs


def preprocess_pdf_text(page_texts: list[str]) -> str:
    page_lines = [[line.strip() for line in text.replace("\r", "\n").splitlines()] for text in page_texts]
    repeated_margin_lines = _detect_repeated_margin_lines(page_lines)

    cleaned_pages: list[str] = []
    for lines in page_lines:
        filtered_lines = [
            line
            for line in lines
            if line.strip()
            and _normalize_margin_line(line) not in repeated_margin_lines
            and not _PAGE_NUMBER_RE.fullmatch(line.strip())
        ]
        paragraphs = _coalesce_pdf_lines(filtered_lines)
        if paragraphs:
            cleaned_pages.append("\n\n".join(paragraphs))

    blocks = [block.strip() for page in cleaned_pages for block in page.split("\n\n") if block.strip()]
    blocks = _trim_pdf_backmatter(blocks)
    text = "\n\n".join(blocks).strip()
    text = re.sub(r"\n{3,}", "\n\n", text)
    if not text:
        raise ValueError("No extractable text after PDF preprocessing")
    return text


def _extract_pdf_opendataloader(path: Path) -> str | None:
    """Extract PDF text via opendataloader-pdf (markdown output)."""
    try:
        import opendataloader_pdf

        with tempfile.TemporaryDirectory() as outdir:
            opendataloader_pdf.convert(
                input_path=[str(path)],
                output_dir=outdir,
                format="markdown",
            )
            md_files = list(Path(outdir).rglob("*.md"))
            if not md_files:
                return None
            text = md_files[0].read_text(encoding="utf-8").strip()
            return text or None
    except Exception:
        logger.warning("opendataloader-pdf failed for %s, falling back to pypdf", path.name, exc_info=True)
        return None


def inspect_pdf(path: Path) -> ParsedPdf:
    reader = PdfReader(path)
    page_count = len(reader.pages)

    text = _extract_pdf_opendataloader(path)
    if text:
        return ParsedPdf(text=text, page_count=page_count)

    # Fall back to pypdf + regex post-processing
    page_texts = [(page.extract_text() or "") for page in reader.pages]
    text = preprocess_pdf_text(page_texts)
    if not text:
        raise ValueError(f"No extractable text in {path.name}")
    return ParsedPdf(text=text, page_count=page_count)


def parse_pdf(path: Path) -> str:
    return inspect_pdf(path).text


def parse_docx(path: Path) -> str:
    doc = DocxDocument(str(path))
    parts: list[str] = []
    # Paragraphs
    for p in doc.paragraphs:
        if p.text.strip():
            parts.append(p.text)
    # Tables (python-docx stores table text separately from paragraphs)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    text = "\n".join(parts).strip()
    if not text:
        raise ValueError(f"No extractable text in {path.name}")
    return text


def parse_markdown(path: Path) -> str:
    text = path.read_text(encoding="utf-8").strip()
    if not text:
        raise ValueError(f"No extractable text in {path.name}")
    return text


_PARSERS = {".pdf": parse_pdf, ".docx": parse_docx, ".md": parse_markdown}


def parse_file(path: Path) -> ParsedFile:
    suffix = path.suffix.lower()
    mime_type = _MIME_MAP.get(suffix)
    if mime_type is None:
        raise UnsupportedFormatError(f"Unsupported format: {suffix}")

    text = _PARSERS[suffix](path)

    return ParsedFile(text=text, mime_type=mime_type, title=path.stem)


async def ingest_file(db: aiosqlite.Connection, path: Path) -> Document:
    parsed = parse_file(path)
    cursor = await db.execute(
        "INSERT INTO documents (source_type, title, text, mime_type, source_uri) "
        "VALUES ('file', ?, ?, ?, ?) RETURNING *",
        (parsed.title, parsed.text, parsed.mime_type, str(path)),
    )
    row = await cursor.fetchone()
    await db.commit()
    assert row is not None
    return Document.from_row(row)
